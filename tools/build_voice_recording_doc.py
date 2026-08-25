from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "voice-recording-script.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 42, 48)
MUTED = RGBColor(105, 119, 127)
GOLD = RGBColor(161, 117, 57)
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF8E8"
LIGHT_GRAY = "F2F4F7"
FONT = "Aptos"
EAST_ASIA_FONT = "Microsoft YaHei"


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size=11, color=INK, bold=False):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_border(paragraph, color="D9E2EA", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), space)
        border.set(qn("w:color"), color)
        borders.append(border)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("ZT.AI · VOICE MODE")
    set_run_font(run, size=10, color=GOLD, bold=True)
    run.font.small_caps = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("三语语音录音稿")
    set_run_font(run, size=28, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("中文 / English / 日本語 · 回家后可以直接照稿录音")
    set_run_font(run, size=13, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("整理日期：2026 年 8 月 25 日")
    set_run_font(run, size=9.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.25
    set_paragraph_border(p, color="E5C98F", size="10", space="8")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GOLD)
    p._p.get_or_add_pPr().append(shading)
    label = p.add_run("录音目标  ")
    set_run_font(label, size=10.5, color=GOLD, bold=True)
    body = p.add_run("共 9 段：中文、英语、日语各 3 段；每段约 60—90 秒。先分别保存原始录音，不要把 9 段一次性合并。")
    set_run_font(body, size=10.5, color=INK)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return paragraph


def add_body(doc, text, *, color=INK, italic=False, after=6, size=11):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, italic=italic)
    return paragraph


def add_note(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_border(paragraph, color="D8E2EA", size="8", space="7")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_BLUE)
    paragraph._p.get_or_add_pPr().append(shading)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=10, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10, color=INK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_script_section(doc, filename, title, note, lines):
    add_heading(doc, f"{filename} · {title}", 2)
    add_note(doc, "录音提示", note)
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(line)
        set_run_font(run, size=11.5, color=INK)


def add_language(doc, heading, subtitle, sections):
    doc.add_page_break()
    add_heading(doc, heading, 1)
    add_body(doc, subtitle, color=MUTED, italic=True, after=12, size=10.5)
    for filename, title, note, lines in sections:
        add_script_section(doc, filename, title, note, lines)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in ((1, 16, BLUE, 18, 10), (2, 13, BLUE, 14, 7), (3, 12, DARK_BLUE, 10, 5)):
        style = doc.styles[f"Heading {level}"]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("ZT.AI · 三语语音录音稿")
    set_run_font(run, size=8.5, color=MUTED)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title(doc)

    add_heading(doc, "一、录音前快速检查", 1)
    add_table(
        doc,
        ["检查项", "建议"],
        [
            ("环境", "安静房间，关闭风扇、空调和音乐；不要有明显回声。"),
            ("距离", "手机或麦克风距离嘴巴约 15—25 厘米，整套录音尽量保持一致。"),
            ("声音", "正常说话，不要故意压低或拔高；不要使用变声、混响、配乐或强美化。"),
            ("格式", "优先 WAV；如果设备不方便，也可以录 M4A 或 MP3。每段单独保存。"),
            ("流程", "每句之间自然停顿约半秒；每段开头和结尾留 1—2 秒安静。"),
        ],
        [1900, 7460],
    )
    add_note(doc, "重要", "MiniMax 官方音色克隆文件要求为 MP3、M4A 或 WAV，单个文件 10 秒至 5 分钟且不超过 20MB。9 段原始录音先分别保存，后续再按语言和效果组合。")

    add_heading(doc, "二、文件命名清单", 1)
    add_table(
        doc,
        ["编号", "文件名", "内容"],
        [
            ("01", "01-zh-natural.wav", "中文：自然介绍"),
            ("02", "02-zh-style.wav", "中文：不同语气"),
            ("03", "03-zh-data.wav", "中文：数字和专有名词"),
            ("04", "04-en-natural.wav", "English: natural introduction"),
            ("05", "05-en-style.wav", "English: varied tone"),
            ("06", "06-en-data.wav", "English: numbers and proper nouns"),
            ("07", "07-ja-natural.wav", "日本語：自然な紹介"),
            ("08", "08-ja-style.wav", "日本語：異なる話し方"),
            ("09", "09-ja-data.wav", "日本語：数字と固有名詞"),
        ],
        [700, 2850, 5810],
    )

    zh = [
        ("01-zh-natural.wav", "中文：自然介绍", "用自然、稳定、清楚的语气；不要像播音员一样过度用力。", [
            "你好，我是蔡宙廷，这里是 ZT.AI。",
            "我们先把问题说清楚，再决定下一步怎么做。",
            "你把目标、限制条件和已有资料告诉我，我会先整理事实，再给你可执行的方案。",
            "如果这个问题涉及当前日期、新闻、价格或者陌生事物，我会先联网核实，不会凭印象乱猜。",
            "我更关注真实结果、成本、效率和风险边界，而不是只看表面热度。",
            "如果证据不够，我会明确告诉你哪里已经确认，哪里还需要验证。",
            "这件事可以做，但要先确认数据、时间和权限是否足够。",
            "好，我们现在开始。",
        ]),
        ("02-zh-style.wav", "中文：不同语气", "依次使用认真、强调重点、温和、稍微积极的语气；不要夸张表演。", [
            "这个结果不对，先别急着下结论，我们把原始数据重新看一遍。",
            "这一步我认可，但还缺一个关键证据。",
            "这个方案可以上线，不过要先把失败时的处理方式设计好。",
            "你说得有道理，我再核实一下具体事实。",
            "这个问题比较复杂，我先给你结论，再解释原因。",
            "如果系统暂时不可用，我不会假装已经完成。",
            "这个版本先保留，后面有新的录音或数据再继续优化。",
            "没关系，我们一步一步来。",
            "谢谢，你提醒得对。",
        ]),
        ("03-zh-data.wav", "中文：数字和专有名词", "数字要读清楚；英文品牌名按你平时的自然发音读，不要刻意变成播音腔。", [
            "今天是二零二六年八月二十五日，现在是北京时间。",
            "版本号是零点二点二六，接口响应时间目标控制在三秒以内。",
            "一百二十三，四千五百六十七点八九，百分之九十九点九。",
            "北京、上海、深圳、东京、伦敦和纽约。",
            "GitHub Pages、Render、Firecrawl、DuckDuckGo 和 MiniMax。",
            "订单号是 ZT-2026-0825-A17。",
            "我们先做一个小范围测试，确认有效以后，再扩大到网页、桌面端和安卓端。",
            "如果搜索失败，就明确说明没有核实到，不能用旧知识补猜。",
        ]),
    ]
    en = [
        ("04-en-natural.wav", "English: Natural Introduction", "Speak naturally and clearly. Do not force an American or British accent.", [
            "Hello, I’m Zhouting Cai. Welcome to ZT.AI.",
            "I will first understand the question, organize the important information, and then give you a clear and practical answer.",
            "If the question involves current events, prices, dates, or unfamiliar topics, I will verify the information online before answering.",
            "I will not make up an answer when the available evidence is incomplete.",
            "I care about real results, cost, efficiency, and the risks behind every decision.",
            "If the problem is complicated, I will explain the conclusion first and then walk through the reasons.",
            "Let’s make the problem clear and decide what to do next.",
        ]),
        ("05-en-style.wav", "English: Varied Tone", "Use a serious tone, a gentle tone, and a slightly positive tone in different sentences.", [
            "This result does not look right. Let’s not jump to a conclusion yet.",
            "I agree with this part, but we still need one important piece of evidence.",
            "The plan is workable, but we need to think about what happens if it fails.",
            "You make a good point. I will check the facts again.",
            "This version can be released, but we should test it before expanding the scope.",
            "If the service is temporarily unavailable, I will say so clearly instead of pretending that the task is complete.",
            "There is no need to rush. We can solve the problem step by step.",
            "Thank you for pointing that out. You are right.",
        ]),
        ("06-en-data.wav", "English: Numbers and Proper Nouns", "Read dates, numbers, product names, and service names slowly enough to be understood.", [
            "Today is August twenty-fifth, twenty twenty-six, and the current time is China Standard Time.",
            "The current version is zero point two point two six.",
            "The response time target is less than three seconds.",
            "The numbers are one hundred and twenty-three, four thousand five hundred and sixty-seven point eight nine, and ninety-nine point nine percent.",
            "The cities are Beijing, Shanghai, Shenzhen, Tokyo, London, and New York.",
            "The services include GitHub Pages, Render, Firecrawl, DuckDuckGo, and MiniMax.",
            "The order number is ZT-2026-0825-A17.",
            "We will start with a small test, confirm that it works, and then expand it to the web, desktop, and Android versions.",
        ]),
    ]
    ja = [
        ("07-ja-natural.wav", "日本語：自然な紹介", "落ち着いた自然な声で、無理にアナウンサーのように読まないでください。", [
            "こんにちは、蔡宙廷です。ZT.AIへようこそ。",
            "まず質問の内容を整理して、大切な情報を確認してから、分かりやすく答えます。",
            "現在のニュース、価格、日付、または知らない内容については、回答する前にインターネットで確認します。",
            "十分な証拠がない場合は、想像で答えることはありません。",
            "私は、実際の結果、コスト、効率、そしてリスクを重視しています。",
            "問題が複雑な場合は、まず結論を説明して、そのあとで理由を説明します。",
            "それでは、問題を整理して、次に何をするか決めましょう。",
        ]),
        ("08-ja-style.wav", "日本語：異なる話し方", "真剣、やさしい、少し前向きな話し方を使い分けてください。大げさに演じる必要はありません。", [
            "この結果は正しくないようです。すぐに結論を出さず、もう一度確認しましょう。",
            "この部分については賛成ですが、まだ大切な証拠が足りません。",
            "この方法は実行できますが、失敗した場合の対応も先に考える必要があります。",
            "なるほど、その意見は分かります。具体的な事実をもう一度確認します。",
            "このバージョンは公開できますが、範囲を広げる前にテストを行いましょう。",
            "サービスが一時的に利用できない場合は、完了したふりをせず、はっきり説明します。",
            "急ぐ必要はありません。一つずつ解決していきましょう。",
            "ご指摘ありがとうございます。その通りです。",
        ]),
        ("09-ja-data.wav", "日本語：数字と固有名詞", "数字、日付、サービス名は、意味が伝わるように少しゆっくり読んでください。", [
            "今日は二〇二六年八月二十五日で、現在の時刻は中国標準時です。",
            "現在のバージョンは、ゼロ点二点二六です。",
            "目標の応答時間は三秒以内です。",
            "数字は、百二十三、四千五百六十七点八九、そして九十九点九パーセントです。",
            "都市は、北京、上海、深圳、東京、ロンドン、そしてニューヨークです。",
            "サービスには、GitHub Pages、Render、Firecrawl、DuckDuckGo、そしてMiniMaxがあります。",
            "注文番号は、ZT-2026-0825-A17です。",
            "まず小さなテストを行い、問題がなければ、ウェブ、デスクトップ、そしてAndroid版へ広げます。",
        ]),
    ]

    add_language(doc, "三、中文录音稿", "中文是核心语音样本，请优先保证这三段的音质、稳定性和自然程度。", zh)
    add_language(doc, "四、English Recording Script", "English is needed for multilingual voice replies. Keep the pronunciation natural and consistent.", en)
    add_language(doc, "五、日本語録音稿", "日本語の音声応答にも使うため、無理のない速さで、言葉をはっきり読んでください。", ja)

    doc.add_page_break()
    add_heading(doc, "六、录完后的交付方式", 1)
    add_body(doc, "请保留 9 个原始文件，不要先压缩、混音或加背景音乐。录完后把 9 个文件放进一个文件夹，或者打包成 ZIP 发给我。", after=8)
    add_body(doc, "如果某一段读错，可以只重录对应文件，不需要全部重录。文件名保持清单中的格式，方便后续检查和组合。", after=8)
    add_note(doc, "提醒", "声音克隆只使用你本人明确授权的声音。之前的视频素材不要直接用于克隆，除非确认里面是你本人的声音，并且你明确同意。")
    add_heading(doc, "官方技术参考", 2)
    add_body(doc, "MiniMax 音色快速复刻：https://platform.minimaxi.com/docs/guides/speech-voice-clone", color=MUTED, size=9.5, after=3)
    add_body(doc, "MiniMax 多语言语音合成：https://platform.minimaxi.com/docs/guides/speech-t2a-async", color=MUTED, size=9.5, after=3)

    doc.core_properties.title = "ZT.AI 三语语音录音稿"
    doc.core_properties.subject = "中文、英语、日语语音克隆录音脚本"
    doc.core_properties.author = "ZT.AI"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
